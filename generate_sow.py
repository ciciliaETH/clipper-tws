from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Style defaults --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

def add_heading_custom(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_bold_paragraph(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_paragraph_with_bold(text):
    """Parse text with **bold** markers"""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle *italic*
            sub_parts = re.split(r'(\*[^*]+\*)', part)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*') and not sp.startswith('**'):
                    run = p.add_run(sp[1:-1])
                    run.italic = True
                else:
                    p.add_run(sp)
    return p

def add_item(letter, bold_text, desc):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f'{letter}. ')
    run2 = p.add_run(bold_text)
    run2.bold = True
    # Parse desc for italic
    parts = re.split(r'(\*[^*]+\*)', desc)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            r = p.add_run(part[1:-1])
            r.italic = True
        else:
            p.add_run(part)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Gray background
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'D9E2F3'
        })
        shading.append(shd)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            is_bold = isinstance(val, str) and val.startswith('**') and val.endswith('**')
            text = val[2:-2] if is_bold else val
            run = p.add_run(str(text))
            run.font.size = Pt(10)
            if is_bold:
                run.bold = True
            # Right-align currency columns
            if isinstance(val, str) and 'Rp' in val:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Center align jam columns
            if c_idx == 2 and len(row) == 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

# ============================================================
# DOCUMENT CONTENT
# ============================================================

# Title
title = doc.add_heading('', level=0)
run = title.add_run('A. Project — Multi-Platform Analytics Dashboard\n(TikTok, Instagram & YouTube)')
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.size = Pt(16)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# --- 1. Project Summary ---
add_heading_custom('1. Project Summary', level=2)

p = doc.add_paragraph()
p.add_run('Developer akan membangun sebuah ')
p.add_run('Multi-Platform Analytics Dashboard').bold = True
p.add_run(' dengan fitur lengkap yang mencakup ')
p.add_run('3 platform sekaligus: TikTok, Instagram, dan YouTube').bold = True
p.add_run('. Dashboard ini meliputi tracking hashtag, trending detection, leaderboard, campaign management system, group analytics, role-based access control, grafik analytics, serta multi-campaign system.')

p2 = doc.add_paragraph(
    'Seluruh platform terintegrasi dalam satu unified dashboard dengan fitur auto-refresh, '
    'historical data tracking, dan full mobile responsive design. Project ini juga mencakup '
    'biaya operasional 1 tahun (API 3 platform + server hosting), sehingga Client tidak perlu '
    'membayar biaya tahunan terpisah selama 12 bulan pertama.'
)

# --- 2. Scope of Work ---
add_heading_custom('2. Scope of Work', level=2)
add_heading_custom('2.1. Development Scope (214 Jam Total)', level=3)

# A. Dashboard & Analytics
add_heading_custom('A. Dashboard & Analytics UI (3 Platform) \u2014 34 Jam', level=4)
doc.add_paragraph(
    'Halaman utama yang menampilkan seluruh data performa konten dari 3 platform dalam satu '
    'tampilan terpadu. Pengguna dapat langsung melihat ringkasan total Views, Likes, Comments, '
    'dan jumlah Posts tanpa perlu membuka masing-masing platform.'
)
add_item('a', 'Dashboard utama (overview)', ' \u2014 halaman ringkasan yang menampilkan total metrik dari seluruh platform sekaligus')
add_item('b', 'Chart analytics lengkap (Line Chart)', ' \u2014 grafik garis interaktif untuk memvisualisasikan tren Views, Likes, dan Comments dari waktu ke waktu')
add_item('c', 'Posts chart per-platform', ' \u2014 grafik terpisah yang menampilkan jumlah konten yang diposting, dengan breakdown warna berbeda untuk TikTok, Instagram, dan YouTube')
add_item('d', 'Historical data integration', ' \u2014 penggabungan data lama (sebelum sistem berjalan) dengan data realtime, sehingga grafik menampilkan data lengkap sejak awal campaign')
add_item('e', 'Multiple chart modes', ' \u2014 dua mode tampilan: *Post Date* (berdasarkan tanggal konten diposting) dan *Accrual* (berdasarkan pertumbuhan metrik harian)')
add_item('f', 'Platform filter', ' \u2014 tombol filter untuk melihat data All / TikTok saja / Instagram saja / YouTube saja')
add_item('g', 'Custom date range', ' \u2014 pengguna dapat memilih rentang tanggal tertentu untuk analisis periode spesifik')
add_item('h', 'Combined header totals', ' \u2014 angka total di header yang menggabungkan data historis + data realtime secara akurat')

# B. TikTok
add_heading_custom('B. TikTok API Integration \u2014 20 Jam', level=4)
doc.add_paragraph(
    'Menghubungkan sistem dengan TikTok untuk mengambil data video (views, likes, comments) dari '
    'setiap akun peserta secara otomatis. Tanpa integrasi ini, data harus diinput manual satu per satu.'
)
add_item('a', 'Setup & autentikasi TikTok API', ' \u2014 konfigurasi koneksi ke TikTok melalui RapidAPI, termasuk manajemen API key dan rate limiting')
add_item('b', 'Daily data fetching', ' \u2014 sistem mengambil data terbaru dari setiap akun TikTok setiap hari secara otomatis')
add_item('c', 'Video metadata extraction', ' \u2014 mengambil detail setiap video: jumlah views, likes, comments, tanggal posting, caption, dan hashtag')
add_item('d', 'Data normalization', ' \u2014 mengolah data mentah dari TikTok ke format standar yang bisa digabungkan dengan platform lain di dashboard')

# C. Instagram
add_heading_custom('C. Instagram Aggregator API Integration \u2014 24 Jam', level=4)
doc.add_paragraph(
    'Menghubungkan sistem dengan Instagram melalui layanan aggregator pihak ketiga (karena Instagram '
    'tidak menyediakan API publik langsung). Proses ini lebih kompleks dari TikTok karena membutuhkan '
    'penanganan khusus untuk resolusi data dan proteksi data.'
)
add_item('a', 'Setup Instagram Aggregator API (v3)', ' \u2014 konfigurasi koneksi ke aggregator server yang menyediakan data Instagram Reels')
add_item('b', 'Daily data fetching', ' \u2014 pengambilan data views, likes, comments dari setiap akun Instagram peserta secara otomatis setiap hari')
add_item('c', 'Data parsing & conversion', ' \u2014 konversi data dari format aggregator (misalnya: timestamp UNIX ke tanggal, shortcode ke numeric ID) agar kompatibel dengan database')
add_item('d', 'NULL protection system', ' \u2014 mekanisme keamanan data: jika API mengembalikan data kosong, sistem tetap mempertahankan data yang sudah tersimpan sebelumnya agar tidak hilang')
add_item('e', 'Backfill system', ' \u2014 sistem recovery untuk mengisi data yang tidak lengkap (misalnya tanggal posting yang belum tersedia saat pengambilan awal)')

# D. YouTube
add_heading_custom('D. YouTube Data API Integration \u2014 20 Jam', level=4)
doc.add_paragraph(
    'Menghubungkan sistem dengan YouTube untuk mengambil data performa video dari channel peserta. '
    'YouTube memiliki API resmi namun memerlukan penanganan quota dan autentikasi khusus.'
)
add_item('a', 'Setup YouTube Data API', ' \u2014 konfigurasi koneksi ke Google/YouTube API, termasuk manajemen API key dan quota')
add_item('b', 'Daily data fetching', ' \u2014 pengambilan data views, likes, comments dari setiap video YouTube peserta secara otomatis')
add_item('c', 'Channel management', ' \u2014 pengelolaan daftar channel YouTube yang dipantau per campaign')
add_item('d', 'Data normalization', ' \u2014 mengolah data YouTube ke format standar yang seragam dengan TikTok dan Instagram di dashboard')

# E. Cron Job
add_heading_custom('E. Cron Job Automation, Retry Queue & Backfill \u2014 16 Jam', level=4)
doc.add_paragraph(
    'Sistem otomatisasi latar belakang yang memastikan data dari 3 platform selalu ter-update tanpa '
    'perlu intervensi manual. Jika terjadi kegagalan (server down, API error), sistem akan mencoba ulang secara otomatis.'
)
add_item('a', 'Cron job scheduling', ' \u2014 penjadwalan otomatis pengambilan data untuk 3 platform (berjalan setiap hari tanpa perlu dijalankan manual)')
add_item('b', 'Retry queue dengan exponential backoff', ' \u2014 jika pengambilan data gagal, sistem menyimpan antrian dan mencoba ulang secara bertahap (1 detik \u2192 2 detik \u2192 4 detik dst.) agar tidak membebani server')
add_item('c', 'Concurrent processing', ' \u2014 pengambilan data beberapa akun sekaligus secara paralel untuk mempercepat proses refresh')
add_item('d', 'Backfill & data recovery', ' \u2014 mengisi ulang data yang hilang atau tidak lengkap dari periode sebelumnya')

# F. Campaign
add_heading_custom('F. Campaign Management System \u2014 22 Jam', level=4)
doc.add_paragraph(
    'Sistem pengelolaan campaign (turnamen/kompetisi) yang memungkinkan admin membuat dan mengelola '
    'beberapa campaign sekaligus, masing-masing dengan peserta, hashtag, dan periode yang berbeda.'
)
add_item('a', 'Multi-campaign CRUD', ' \u2014 membuat, mengedit, dan menghapus campaign dengan pengaturan nama, hashtag, dan periode')
add_item('b', 'Participant management', ' \u2014 mengelola daftar peserta per campaign untuk masing-masing platform (akun TikTok, Instagram, YouTube)')
add_item('c', 'Campaign hashtag tracking', ' \u2014 melacak konten berdasarkan hashtag campaign di semua platform')
add_item('d', 'Campaign videos detail page', ' \u2014 halaman yang menampilkan semua video yang terdeteksi dalam sebuah campaign, lengkap dengan metrik per video')
add_item('e', 'Per-campaign analytics', ' \u2014 grafik dan statistik khusus per campaign, sehingga bisa membandingkan performa antar campaign')

# G. Group & Leaderboard
add_heading_custom('G. Group & Leaderboard System \u2014 22 Jam', level=4)
doc.add_paragraph(
    'Sistem pengelompokan peserta dan peringkat untuk memantau performa individu maupun tim. '
    'Berguna untuk melihat siapa peserta terbaik dan video mana yang paling viral.'
)
add_item('a', 'Group management', ' \u2014 membuat grup/tim, menambah dan mengelola anggota, serta menentukan struktur tim')
add_item('b', 'Group analytics dashboard', ' \u2014 grafik dan statistik per grup, termasuk perbandingan performa antar anggota dalam satu tampilan')
add_item('c', 'Cross-platform leaderboard', ' \u2014 papan peringkat yang menggabungkan data dari 3 platform, diurutkan berdasarkan total views, likes, atau comments')
add_item('d', 'Employee detail page', ' \u2014 halaman profil per peserta yang menampilkan semua video dari seluruh platform beserta riwayat performanya')
add_item('e', 'Top viral videos detection', ' \u2014 deteksi otomatis video dengan performa tertinggi (views terbanyak, engagement tertinggi) dari seluruh platform')

# H. Admin Panel
add_heading_custom('H. Admin Panel \u2014 18 Jam', level=4)
doc.add_paragraph(
    'Panel kontrol khusus admin untuk mengelola seluruh sistem, termasuk memicu refresh data secara '
    'manual, memonitor status API, dan mengelola hadiah/reward campaign.'
)
add_item('a', 'Admin dashboard', ' \u2014 halaman overview khusus admin dengan ringkasan semua campaign yang berjalan')
add_item('b', 'Auto-refresh 3 platform', ' \u2014 tombol untuk memperbarui data dari TikTok, Instagram, dan YouTube secara bersamaan atau per platform')
add_item('c', 'Status monitoring', ' \u2014 indikator real-time yang menampilkan status terakhir refresh (berhasil/gagal, jumlah data yang diambil, waktu terakhir update)')
add_item('d', 'Retry queue UI', ' \u2014 tampilan antrian data yang gagal diambil, dengan opsi retry manual')
add_item('e', 'Prize/reward management', ' \u2014 pengelolaan hadiah per campaign (input hadiah, tampilkan di dashboard peserta)')

# I. Auth
add_heading_custom('I. User Authentication & Access Control \u2014 8 Jam', level=4)
doc.add_paragraph(
    'Sistem login dan kontrol akses untuk memastikan hanya pengguna yang berwenang yang dapat mengakses '
    'fitur tertentu. Admin memiliki akses penuh, sementara User/Public hanya dapat melihat data.'
)
add_item('a', 'Login & authentication', ' \u2014 halaman login yang aman dengan manajemen session')
add_item('b', 'Role-Based Access Control (RBAC)', ' \u2014 pembagian hak akses: Admin (akses penuh, kelola campaign, refresh data) dan User/Public (hanya lihat dashboard & leaderboard)')

# J. Mobile
add_heading_custom('J. Mobile Responsive Design \u2014 14 Jam', level=4)
doc.add_paragraph(
    'Penyesuaian tampilan seluruh halaman agar dapat diakses dengan nyaman di perangkat mobile '
    '(HP/tablet), tanpa perlu membuat aplikasi terpisah.'
)
add_item('a', 'Responsive dashboard & analytics', ' \u2014 grafik, tabel, dan filter yang menyesuaikan ukuran layar HP')
add_item('b', 'Responsive admin panel', ' \u2014 panel admin yang tetap fungsional di layar kecil (tombol, modal, tabel)')
add_item('c', 'Responsive leaderboard & groups', ' \u2014 tampilan peringkat dan grup yang optimal di mobile')
add_item('d', 'Responsive detail pages', ' \u2014 halaman detail employee, video, dan campaign yang mobile-friendly')

# K. DB & Deploy
add_heading_custom('K. Database Design, Deployment & QA \u2014 16 Jam', level=4)
doc.add_paragraph(
    'Perancangan struktur database, deployment ke server production, dan pengujian menyeluruh untuk '
    'memastikan sistem berjalan stabil dan tanpa error.'
)
add_item('a', 'Database schema design (Supabase)', ' \u2014 perancangan tabel, relasi, dan index untuk menyimpan data 3 platform secara efisien')
add_item('b', 'Server setup & deployment (Vercel)', ' \u2014 konfigurasi server production, domain, environment variables, dan CI/CD')
add_item('c', 'Cross-platform testing', ' \u2014 pengujian fungsi pengambilan data dari 3 platform untuk memastikan akurasi')
add_item('d', 'Performance optimization', ' \u2014 optimasi kecepatan loading dashboard, query database, dan API calls')
add_item('e', 'Dokumentasi penggunaan', ' \u2014 panduan cara menggunakan dashboard, admin panel, dan fitur-fitur utama')

# 2.2 Out of Scope
add_heading_custom('2.2. Tidak Termasuk (Out of Scope)', level=3)
doc.add_paragraph('Fitur di bawah tidak termasuk dalam versi ini:')
for item in ['Discord Bot integration', 'Export data CSV & PDF', 'Advanced filtering (AI, rekomendasi otomatis, scoring cerdas)', 'Pembuatan mobile app (Android/iOS)']:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

# --- 3. Deliverables ---
add_heading_custom('3. Deliverables', level=2)
doc.add_paragraph('Developer akan menyerahkan:')
deliverables = [
    'Multi-Platform Analytics Dashboard live (web) \u2014 TikTok, Instagram, YouTube',
    'Backend + API + cron job berjalan untuk 3 platform',
    'Campaign Management System (multi-platform)',
    'Group & Leaderboard system (cross-platform)',
    'Admin Panel dengan kontrol refresh 3 platform',
    'Trending detection engine & top viral videos',
    'Historical data system (weekly historical + realtime merge)',
    'Full mobile responsive design',
    'Role-based access control',
    'Dokumentasi penggunaan',
    '30 hari support pasca-handover',
]
for i, d in enumerate(deliverables):
    letter = chr(ord('a') + i)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{letter}. {d}')

# --- 4. Biaya ---
add_heading_custom('4. Biaya & Struktur Pembayaran', level=2)

add_heading_custom('4.1. Rincian Biaya Development (Rate: Rp 157.500/jam)', level=3)

dev_rows = [
    ['A', 'Dashboard & Analytics UI (3 Platform)', '34', 'Rp 5.355.000'],
    ['B', 'TikTok API Integration', '20', 'Rp 3.150.000'],
    ['C', 'Instagram Aggregator API Integration', '24', 'Rp 3.780.000'],
    ['D', 'YouTube Data API Integration', '20', 'Rp 3.150.000'],
    ['E', 'Cron Job Automation, Retry Queue & Backfill', '16', 'Rp 2.520.000'],
    ['F', 'Campaign Management System', '22', 'Rp 3.465.000'],
    ['G', 'Group & Leaderboard System', '22', 'Rp 3.465.000'],
    ['H', 'Admin Panel & Platform Controls', '18', 'Rp 2.835.000'],
    ['I', 'User Authentication & RBAC', '8', 'Rp 1.260.000'],
    ['J', 'Mobile Responsive Design', '14', 'Rp 2.205.000'],
    ['K', 'Database Design, Deployment & QA', '16', 'Rp 2.520.000'],
    ['', '**Subtotal Development (214 jam)**', '**214**', '**Rp 33.705.000**'],
]
add_table(['No', 'Modul (sesuai Scope 2.1)', 'Jam', 'Biaya'], dev_rows, [1.2, 9, 1.5, 3.5])

doc.add_paragraph('')  # spacing

add_heading_custom('4.2. Biaya Operasional 1 Tahun (Sudah Termasuk Total)', level=3)

ops_rows = [
    ['1', 'TikTok API (RapidAPI subscription)', 'Rp 3.000.000'],
    ['2', 'Instagram Aggregator API (v3 license)', 'Rp 2.500.000'],
    ['3', 'YouTube Data API (quota extension)', 'Rp 1.500.000'],
    ['4', 'Virtual Private Server (hosting & scaling)', 'Rp 2.393.200'],
    ['', '**Subtotal Operasional**', '**Rp 9.393.200**'],
]
add_table(['No', 'Item', 'Biaya / Tahun'], ops_rows, [1.2, 9, 3.5])

doc.add_paragraph('')

add_heading_custom('4.3. Total Biaya Project', level=3)

total_rows = [
    ['Development (214 jam)', 'Rp 33.705.000'],
    ['Operasional 1 tahun (API 3 Platform + Server)', 'Rp 9.393.200'],
    ['**TOTAL PROJECT**', '**Rp 43.098.200**'],
]
add_table(['Kategori', 'Biaya'], total_rows, [9, 4])

doc.add_paragraph('')

add_heading_custom('4.4. Payment Terms', level=3)
p = doc.add_paragraph(style='List Bullet')
p.add_run('50% DP di awal \u2014 ')
p.add_run('Rp 21.549.100').bold = True

p = doc.add_paragraph(style='List Bullet')
p.add_run('50% setelah project selesai & deployed \u2014 ')
p.add_run('Rp 21.549.100').bold = True

# --- 5. Timeline ---
add_heading_custom('5. Timeline Pengerjaan (3 Bulan)', level=2)

timeline_rows = [
    ['1', 'Design & System Architecture', 'Sprint 1', 'Finalisasi UI/UX + perancangan flow & database'],
    ['2', 'TikTok API Integration', 'Sprint 2-3', 'Setup API, autentikasi, daily tracking'],
    ['3', 'Instagram API Integration', 'Sprint 4-5', 'Setup aggregator API, daily tracking, retry queue'],
    ['4', 'YouTube API Integration', 'Sprint 6-7', 'Setup YouTube API, daily tracking, auto-refresh'],
    ['5', 'Dashboard Multi-Platform & Analytics', 'Sprint 8-9', 'Unified analytics, chart, historical data merge'],
    ['6', 'Campaign, Group & Leaderboard', 'Sprint 10', 'Multi-platform campaign, cross-platform leaderboard'],
    ['7', 'Admin Panel & Mobile Responsive', 'Sprint 11', 'Admin controls 3 platform, full mobile responsive'],
    ['8', 'Quality Assurance & Deployment', 'Sprint 12', 'Testing lengkap, performance optimization & deploy'],
]
add_table(['No', 'Deliverable', 'Sprint', 'Deskripsi'], timeline_rows, [1, 5, 2, 7])

# --- 6. Support ---
add_heading_custom('6. Support & Maintenance', level=2)
doc.add_paragraph('Developer menyediakan:')
p = doc.add_paragraph(style='List Bullet')
p.add_run('30 hari support gratis (bugfix minor) pasca-handover')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Perubahan besar / fitur baru dikenakan biaya tambahan')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Maintenance bulanan (opsional): ').bold = True
p.add_run('Rp 3.000.000/bulan \u2014 mencakup monitoring, bug fix, update API jika ada perubahan dari platform, serta minor enhancement')

# --- 7. HKI ---
add_heading_custom('7. Hak Kekayaan Intelektual (HKI)', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Semua source code dan aset menjadi milik Client setelah pembayaran lunas.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Developer dapat menyimpan salinan untuk portofolio privat (tanpa data sensitif).')

# --- D. Persetujuan ---
doc.add_page_break()
add_heading_custom('D. Persetujuan', level=2)
doc.add_paragraph(
    'Dengan menandatangani dokumen ini, kedua belah pihak menyetujui seluruh '
    'ketentuan dalam Statement of Work ini.'
)

doc.add_paragraph('')
p = doc.add_paragraph('Tangerang, __ ___________ 2026')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
doc.add_paragraph('')

# Signature table (no borders)
sig_table = doc.add_table(rows=4, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Remove borders
for row in sig_table.rows:
    for cell in row.cells:
        tc = cell._element.get_or_add_tcPr()
        borders = tc.makeelement(qn('w:tcBorders'), {})
        for border_name in ['top', 'left', 'bottom', 'right']:
            b = borders.makeelement(qn(f'w:{border_name}'), {qn('w:val'): 'none', qn('w:sz'): '0'})
            borders.append(b)
        tc.append(borders)

sig_table.rows[0].cells[0].paragraphs[0].add_run('Pihak Client').bold = True
sig_table.rows[0].cells[1].paragraphs[0].add_run('Pihak Developer').bold = True
sig_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

sig_table.rows[1].cells[0].paragraphs[0].add_run('Trade With Suli')

# Empty space for signature
sig_table.rows[2].cells[0].paragraphs[0].add_run('')
sig_table.rows[2].cells[1].paragraphs[0].add_run('')

sig_table.rows[3].cells[0].paragraphs[0].add_run('Andrew Jonathan').bold = True
r = sig_table.rows[3].cells[1].paragraphs[0].add_run('Naufal Ahmad Fadillah')
r.bold = True
sig_table.rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Save
output_path = r'c:\nofall\dashboard-clipper-V2-main - Copy\dashboard-clipper-V2-main\.claude\worktrees\exciting-tu\SOW - Multi-Platform Analytics Dashboard (TikTok, Instagram & YouTube).docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
